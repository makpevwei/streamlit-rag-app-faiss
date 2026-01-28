import warnings
import os
import sys

# Suppress warnings and set threading for macOS stability
warnings.filterwarnings("ignore")
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["OPENBLAS_NUM_THREADS"] = "1"
os.environ["VECLIB_MAXIMUM_THREADS"] = "1"
os.environ["NUMEXPR_NUM_THREADS"] = "1"

from multiprocessing import resource_tracker

warnings.filterwarnings(
    "ignore",
    message="resource_tracker: There appear to be [0-9]+ leaked semaphore objects",
)

import os
import bs4
import streamlit as st
import numpy as np
import pandas as pd
import faiss
from pathlib import Path
from io import BytesIO 
from pypdf import PdfReader
from docx import Document
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from typing import List
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings


# --- 1. Environment & API Setup ---
os.environ["TOKENIZERS_PARALLELISM"] = "false"
load_dotenv()

groq_api_key = os.environ.get("GROQ_API_KEY")
groq_model = os.environ.get("GROQ_MODEL") or "llama-3.1-8b-instant"

# --- Configuration ---

EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHAT_MODEL = "llama-3.1-8b-instant"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 300

# Supported File Types
SUPPORTED_TYPES = ["txt","md","pdf","docx","html","csv","xlsx"]


def get_api_key() -> str | None:
    """Get API key from secrets.toml or environment."""
    # Try Streamlit secrets first
    if "GROQ_API_KEY" in st.secrets:
        return st.secrets["GROQ_API_KEY"]
    # Fallback to environment variable
    if os.getenv("GROQ_API_KEY"):
        return os.getenv("GROQ_API_KEY")
    return None


# Document Readers
def read_txt(file) -> str:
    """Read plain text or markdown file."""
    return file.read().decode("utf-8")


def read_pdf(file) -> str:
    """Extract text from pdf files."""
    reader = PdfReader(file)
    text_parts = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"[page {i+1}]\n{page_text}")  
              
    return "\n\n".join(text_parts)


def read_docx(file) -> str:
    """Extract text from Word documents."""
    doc = Document(file)
    paragraphs = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
            
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
                
    return "\n\n".join(paragraphs)


def read_csv(file) -> str:
    """Extract text from CSV files."""
    df = pd.read_csv(file)
    
    # Create a text representation
    text_parts = [f"CSV Data with {len(df)} rows and {len(df.columns)} columns."]
    text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")   
    text_parts.append("\nData:\n") 
    text_parts.append(df.to_markdown(index=False))
     
    return "\n".join(text_parts)


def read_excel(file) -> str:
    """Extract text from Excel files."""
    xlsx = pd.ExcelFile(file)
    text_parts = []
    
    for sheet_name in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet_name)
        text_parts.append(f"\n###Sheet: {sheet_name}")
        text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")   
        text_parts.append(df.to_markdown(index=False))
        
    return "\n\n".join(text_parts)



def read_html(file) -> str:
    """Extract text from HTML files."""
    content = file.read()
    soup = BeautifulSoup(content, 'html.parser')
    
    # Remove scripts and style elements
    for script_or_style in soup(["script", "style", "nav", "footer", "header"]):
        script_or_style.decompose()
    
    # Get text with some structure preserved
    text = soup.get_text(separator="\n")
    
    # Clean up excessive whitespace
    lines = (line.strip() for line in text.splitlines() if line.strip())
    return '\n'.join(lines)



def read_file(file) -> tuple[str, str]:
    """
    Read file based on its extension.
    Returns tuple of (content, file_type)
    """
    filename = file.name.lower()
    
    if filename.endswith(".pdf"):
        return read_pdf(file), "PDF"
    elif filename.endswith(".csv"):
        return read_csv(file), "CSV"
    elif filename.endswith(".docx"):
        return read_docx(file), "WORD"
    elif filename.endswith(".xlsx"):
        return read_excel(file), "EXCEL"
    elif filename.endswith(".html") or filename.endswith(".htm"):
        return read_html(file), "HTML"
    else: # txt, md or unknown text
        return read_txt(file), "TEXT"
    

def chunk_text(
    text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP
) -> list[str]:
    """Split text into overlapping chunks."""
    words = text.split()
    chunks = []

    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)

    return chunks


# Create global embeddings object once - lazy loaded with caching
@st.cache_resource
def load_embeddings():
    """Load embeddings model with proper device handling."""
    import os
    os.environ["TOKENIZERS_PARALLELISM"] = "false"
    os.environ["OMP_NUM_THREADS"] = "1"
    os.environ["MKL_NUM_THREADS"] = "1"
    
    try:
        # For macOS, force CPU and disable threading
        import torch
        torch.set_num_threads(1)
        device = "cpu"  # macOS has issues with GPU, use CPU
    except:
        device = "cpu"
    
    try:
        return HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={"device": device, "trust_remote_code": True},
            encode_kwargs={"normalize_embeddings": True},
        )
    except Exception as e:
        st.error(f"Error loading embeddings: {e}")
        raise

hf_embeddings = load_embeddings()


def get_embeddings(texts: List[str], api_key: str | None = None) -> np.ndarray:
    """Get embeddings using a Hugging Face embeddings model (api_key ignored)."""
    vectors = hf_embeddings.embed_documents(texts)  # List[List[float]]
    return np.array(vectors, dtype="float32")


def create_faiss_index(embeddings: np.ndarray) -> faiss.IndexFlatL2:
    """Create a FAISS index from embeddings."""
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    return index



def search_similar(
    query: str,
    index: faiss.IndexFlatL2,
    chunks: List[str],
    api_key: str | None = None,
    top_k: int = 3,
) -> List[str]:
    """Find most similar chunks to the query."""
    query_embedding = get_embeddings([query], api_key)  # shape (1, dim)
    distances, indices = index.search(query_embedding, top_k)

    results = [chunks[i] for i in indices[0] if i < len(chunks)]
    return results


# Initialize Groq chat model once
groq_chat = ChatGroq(
    model_name=CHAT_MODEL,
    groq_api_key=groq_api_key,  # from your environment
    temperature=0.3
)


def generate_answer(query: str, context_chunks: List[str]) -> str:
    """Generate answer using Groq chat completion with retrieved context."""
    context = "\n\n---\n\n".join(context_chunks)

    system_prompt = """You are a helpful assistant that answers questions based on the provided context.
    
Rules:
- Answer ONLY based on the context provided
- If the context doesn't contain relevant information, say "I couldn't find relevant information in the documents."
- Be concise and direct
- Quote relevant parts when appropriate
- If the question asks for data from tables/spreadsheets, present it clearly"""

    user_prompt = f"""Context:
{context}

---

Question: {query}

Answer:"""

    # Using LangChain-style invocation
    messages = [
        ("system", system_prompt),
        ("user", user_prompt),
    ]

    response = groq_chat.invoke(messages)
    return response.content



# Streamlit UI

def main():
    st.set_page_config(page_title="RAG Q&A System", page_icon="📚", layout="wide")

    st.title("📚 RAG Document Q&A System")
    st.markdown("*Upload documents, ask questions, get AI-powered answers*")

    # Initialize session state
    if "chunks" not in st.session_state:
        st.session_state.chunks = []
    if "index" not in st.session_state:
        st.session_state.index = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "documents_loaded" not in st.session_state:
        st.session_state.documents_loaded = False
    if "doc_stats" not in st.session_state:
        st.session_state.doc_stats = []

    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")

        # Get API key from secrets
        api_key = get_api_key()

        if api_key:
            st.success("✅ API Key loaded")
        else:
            st.error("❌ API Key not found")
            st.markdown(
                """
            **Setup required:**
            
            Create `.streamlit/secrets.toml`:
            ```toml
            GROQ_API_KEY = "your-key-here"
            ```
            
            Or set environment variable:
            ```bash
            export GROQ_API_KEY="your-key"
            ```
            """
            )

        st.divider()

        # Document upload section
        st.header("📄 Add Documents")

        # Show supported formats
        with st.expander("Supported Formats"):
            st.markdown(
                """
            | Format | Extension |
            |--------|-----------|
            | Plain Text | `.txt` |
            | Markdown | `.md` |
            | PDF | `.pdf` |
            | Word | `.docx` |
            | CSV | `.csv` |
            | Excel | `.xlsx` |
            | HTML | `.html` |
            """
            )

        # File upload
        uploaded_files = st.file_uploader(
            "Upload files",
            type=SUPPORTED_TYPES,
            accept_multiple_files=True,
            help="Upload one or more documents",
        )

        # Or paste text
        pasted_text = st.text_area(
            "Or paste text directly", height=120, placeholder="Paste content here..."
        )

        # Advanced settings
        with st.expander("Advanced Settings"):
            chunk_size = st.slider("Chunk size (words)", 100, 1000, CHUNK_SIZE, 50)
            chunk_overlap = st.slider(
                "Chunk overlap (words)", 0, 200, CHUNK_OVERLAP, 10
            )
            top_k = st.slider("Results to retrieve", 1, 10, 3)

        # Process button
        if st.button("🔄 Process Documents", type="primary", disabled=not api_key):
            all_text = ""
            doc_stats = []

            # Process uploaded files
            if uploaded_files:
                progress = st.progress(0, "Processing files...")

                for i, file in enumerate(uploaded_files):
                    try:
                        content, file_type = read_file(file)
                        word_count = len(content.split())
                        all_text += f"\n\n--- Document: {file.name} ---\n\n{content}"
                        doc_stats.append(
                            {"name": file.name, "type": file_type, "words": word_count}
                        )
                        progress.progress(
                            (i + 1) / len(uploaded_files), f"Processed {file.name}"
                        )
                    except Exception as e:
                        st.error(f"Error reading {file.name}: {str(e)}")

                progress.empty()

            # Add pasted text
            if pasted_text.strip():
                all_text += f"\n\n--- Pasted Text ---\n\n{pasted_text}"
                doc_stats.append(
                    {
                        "name": "Pasted text",
                        "type": "Text",
                        "words": len(pasted_text.split()),
                    }
                )

            if all_text.strip():
                with st.spinner("Creating embeddings and index..."):
                    try:
                        # Chunk the text
                        chunks = chunk_text(all_text, chunk_size, chunk_overlap)
                        st.session_state.chunks = chunks

                        # Create embeddings and index
                        embeddings = get_embeddings(chunks, api_key)
                        st.session_state.index = create_faiss_index(embeddings)
                        st.session_state.documents_loaded = True
                        st.session_state.doc_stats = doc_stats
                        st.session_state.top_k = top_k

                        st.success(f"✅ Processed {len(chunks)} chunks!")

                    except Exception as e:
                        st.error(f"Error: {str(e)}")
            else:
                st.warning("Please upload files or paste text")

        # Show loaded documents
        if st.session_state.documents_loaded:
            st.divider()
            st.subheader("📊 Loaded Documents")

            for doc in st.session_state.doc_stats:
                st.markdown(
                    f"**{doc['name']}**  \n{doc['type']} • {doc['words']:,} words"
                )

            st.markdown(f"**Total chunks:** {len(st.session_state.chunks)}")

        # Clear button
        if st.session_state.documents_loaded:
            st.divider()
            if st.button("🗑️ Clear All"):
                st.session_state.chunks = []
                st.session_state.index = None
                st.session_state.chat_history = []
                st.session_state.documents_loaded = False
                st.session_state.doc_stats = []
                st.rerun()

    # Main chat interface

    if not api_key:
        st.warning(
            "⚠️ Groq API key not configured. See sidebar for setup instructions."
        )
        return

    if not st.session_state.documents_loaded:
        st.info("👈 Upload documents in the sidebar, then click 'Process Documents'")

        # Show example
        with st.expander("📖 How it works"):
            st.markdown(
                """
            **RAG (Retrieval-Augmented Generation) Process:**
            
            1. **Upload** — Add your documents (PDF, Word, Excel, CSV, Text, HTML)
            2. **Process** — Documents are split into chunks and converted to embeddings
            3. **Index** — Embeddings are stored in a FAISS vector index
            4. **Query** — Your questions are matched against the document chunks
            5. **Generate** — Relevant chunks are sent to Groq to generate accurate answers
            
            **Supported document types:**
            - 📄 PDF documents
            - 📝 Word documents (.docx)
            - 📊 Spreadsheets (CSV, Excel)
            - 🌐 HTML files
            - 📃 Text and Markdown files
            """
            )
        return

    # Display chat history
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if "sources" in msg and msg["sources"]:
                with st.expander("📎 Sources"):
                    for i, source in enumerate(msg["sources"], 1):
                        st.markdown(f"**Chunk {i}:**")
                        display_text = (
                            source[:500] + "..." if len(source) > 500 else source
                        )
                        st.markdown(f"> {display_text}")

    # Chat input
    if query := st.chat_input("Ask a question about your documents..."):
        # Add user message
        st.session_state.chat_history.append({"role": "user", "content": query})

        with st.chat_message("user"):
            st.markdown(query)

        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("Searching and generating answer..."):
                try:
                    # Get top_k from session or default
                    top_k = st.session_state.get("top_k", 3)

                    # Search for relevant chunks
                    relevant_chunks = search_similar(
                        query,
                        st.session_state.index,
                        st.session_state.chunks,
                        api_key,
                        top_k=top_k,
                    )

                    # Generate answer
                    answer = generate_answer(query, relevant_chunks)

                    st.markdown(answer)

                    # Show sources
                    with st.expander("📎 Sources"):
                        for i, source in enumerate(relevant_chunks, 1):
                            st.markdown(f"**Chunk {i}:**")
                            display_text = (
                                source[:500] + "..." if len(source) > 500 else source
                            )
                            st.markdown(f"> {display_text}")

                    # Save to history
                    st.session_state.chat_history.append(
                        {
                            "role": "assistant",
                            "content": answer,
                            "sources": relevant_chunks,
                        }
                    )

                except Exception as e:
                    st.error(f"Error: {str(e)}")


if __name__ == "__main__":
    main()
